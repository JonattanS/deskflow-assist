import pandas as pd
from docx import Document
import os
from datetime import datetime

# Rutas
ONEDRIVE_PATH = r"C:\Users\MCA칌AS\OneDrive - Nova Corp SAS"
EXCEL_PATH = os.path.join(ONEDRIVE_PATH, "Documentos", "clientes.xlsx")
TEMPLATE_PATH1 = "PROPUESTA FE Y NE RENOVACION 202X -CLIENTE1.docx"
TEMPLATE_PATH2 = "PROPUESTA FE RENOVACION 202X -CLIENTE2.docx"
OUTPUT_DIR = os.path.join(ONEDRIVE_PATH, "Documentos_Generados", "Renovaciones")

# Crear carpeta de salida si no existe
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Cargar Excel
df = pd.read_excel(EXCEL_PATH)

# Funci칩n para reemplazar etiquetas en p치rrafos y tablas
def reemplazar_etiquetas(doc, datos):
    fecha_actual = datetime.now()
    meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
             'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
    
    datos_fecha = {
        'd칤a': str(fecha_actual.day),
        'Mes': meses[fecha_actual.month - 1],
        'A칌O': str(fecha_actual.year),  # A칌O en may칰scula porque en la plantilla lo est치 as칤
        'MES': meses[fecha_actual.month - 1].upper(),  # para el <MES> del pie de p치gina si se escribe en may칰scula
        'a침o': str(fecha_actual.year),
        'mes': meses[fecha_actual.month - 1]
    }

    if 'Nombres y Apellidos' in datos:
        datos['Nombre'] = str(datos['Nombres y Apellidos']).split()[0]
    
    todos_datos = {**datos_fecha, **datos}

    def reemplazar_en_parrafos(parrafos):
        for p in parrafos:
            for key, value in todos_datos.items():
                etiqueta1 = f"<{key}>"
                etiqueta2 = f"춺 {key}췉"
                if etiqueta1 in p.text or etiqueta2 in p.text:
                    for run in p.runs:
                        run.text = run.text.replace(etiqueta1, str(value))
                        run.text = run.text.replace(etiqueta2, str(value))

    def reemplazar_en_tablas(tablas):
        for table in tablas:
            for row in table.rows:
                for cell in row.cells:
                    reemplazar_en_parrafos(cell.paragraphs)

    reemplazar_en_parrafos(doc.paragraphs)
    reemplazar_en_tablas(doc.tables)

    # 游 Encabezado y pie de p치gina
    for section in doc.sections:
        reemplazar_en_parrafos(section.footer.paragraphs)
        reemplazar_en_parrafos(section.header.paragraphs)


# Generar documentos
for idx, row in df.iterrows():
    nit = str(row.iloc[0])
    indicador = row.get('Indicador Tarifa', '')

    # Elegir plantilla seg칰n el valor num칠rico del indicador
    if pd.isna(indicador):
        print(f" Indicador vac칤o para NIT {nit}. Se usar치 plantilla 1 por defecto.")
        plantilla = TEMPLATE_PATH1
    elif str(indicador).strip() == "1":
        plantilla = TEMPLATE_PATH1
    elif str(indicador).strip() == "2":
        plantilla = TEMPLATE_PATH2
    else:
        print(f" Indicador no reconocido ('{indicador}') para NIT {nit}. Se usar치 plantilla 1 por defecto.")
        plantilla = TEMPLATE_PATH1

    # Cargar plantilla seleccionada
    doc = Document(plantilla)

    # Crear carpeta para cliente
    cliente_dir = os.path.join(OUTPUT_DIR, nit)
    os.makedirs(cliente_dir, exist_ok=True)

    # Agregar mapeo espec칤fico para RAZ칍N SOCIAL
    datos_cliente = row.to_dict()
    if 'RAZ칍N SOCIAL' in datos_cliente or 'Raz칩n Social' in datos_cliente:
        # Usar el campo tal como est치 en el Excel
        pass
    elif 'Razon Social' in datos_cliente:
        datos_cliente['RAZ칍N SOCIAL'] = datos_cliente['Razon Social']
    
    # Reemplazar etiquetas
    reemplazar_etiquetas(doc, datos_cliente)

    # Guardar documento
    nombre_archivo = os.path.join(cliente_dir, f"Renovacion_{datetime.now().year}_{nit}.docx")
    doc.save(nombre_archivo)
    print(f" Documento generado: {nombre_archivo}")
