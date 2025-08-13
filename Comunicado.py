import pandas as pd
from docx import Document
import os
from datetime import datetime

# Rutas
ONEDRIVE_PATH = r"C:\Users\MCAÑAS\OneDrive - Nova Corp SAS"
EXCEL_PATH = os.path.join(ONEDRIVE_PATH, "Documentos", "clientes.xlsx")
TEMPLATE_PATH = "RENOVACION_202X_CLIENTE.docx"

OUTPUT_DIR = os.path.join(ONEDRIVE_PATH, "Documentos_Generados", "Comunicados")

# Crear carpeta de salida si no existe
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Cargar Excel
df = pd.read_excel(EXCEL_PATH)

# Leer plantilla Word
with open(TEMPLATE_PATH, 'rb') as file:
    template_doc = Document(file)

# Función para reemplazar etiquetas preservando formato
def reemplazar_etiquetas(doc, datos):
    # Obtener fecha actual
    fecha_actual = datetime.now()
    meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
             'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
    
    # Crear diccionario con datos de fecha
    datos_fecha = {
        'día': str(fecha_actual.day),
        'Mes': meses[fecha_actual.month - 1],
        'año': str(fecha_actual.year)
    }
    
    # Extraer solo el primer nombre si existe el campo "Nombres y Apellidos"
    if 'Nombres y Apellidos' in datos:
        primer_nombre = str(datos['Nombres y Apellidos']).split()[0]
        datos['Nombre'] = primer_nombre
    
    # Combinar datos de fecha con datos del Excel
    todos_datos = {**datos_fecha, **datos}
    
    # Función auxiliar para reemplazar en párrafos
    def reemplazar_en_parrafos(parrafos):
        for p in parrafos:
            texto_completo = p.text
            for key, value in todos_datos.items():
                etiqueta_angular = f"<{key}>"
                etiqueta_guillemet = f"« {key}»"
                
                if etiqueta_angular in texto_completo or etiqueta_guillemet in texto_completo:
                    # Reemplazar en el texto completo del párrafo
                    nuevo_texto = texto_completo.replace(etiqueta_angular, str(value))
                    nuevo_texto = nuevo_texto.replace(etiqueta_guillemet, str(value))
                    
                    # Si hubo cambios, limpiar el párrafo y agregar el nuevo texto
                    if nuevo_texto != texto_completo:
                        p.clear()
                        p.add_run(nuevo_texto)
                        texto_completo = nuevo_texto
    
    # Reemplazar en párrafos principales
    reemplazar_en_parrafos(doc.paragraphs)
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_en_parrafos(cell.paragraphs)

# Generar un documento por cada fila del Excel
for idx, row in df.iterrows():
    # Obtener NIT (primera columna)
    nit = str(row.iloc[0])  # Primera columna del Excel
    
    # Crear carpeta para el cliente usando el NIT
    cliente_dir = os.path.join(OUTPUT_DIR, nit)
    os.makedirs(cliente_dir, exist_ok=True)
    
    # Cargar nuevamente la plantilla para cada cliente
    doc = Document(TEMPLATE_PATH)
    reemplazar_etiquetas(doc, row.to_dict())

    # Guardar documento en la carpeta del cliente
    nombre_archivo = os.path.join(cliente_dir, f"Comunicado_"+str(datetime.now().year)+"_"+nit+".docx")
    doc.save(nombre_archivo)
    print(f" Documento generado: {nombre_archivo}")
